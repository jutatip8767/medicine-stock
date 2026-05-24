import streamlit as st
import pandas as pd
from pathlib import Path
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="ตัดยอดยา", layout="wide")
st.title("โปรแกรมตัดยอดยาในสถานพยาบาล")

DATA_FILE = Path("medicine_stock.csv")


def create_template() -> pd.DataFrame:
    return pd.DataFrame(
        columns=[
            "วันที่รับยา",
            "ชื่อยา",
            "ยอดรับเข้า",
            "ยอดใช้ไป",
            "ราคาต่อหน่วย",
        ]
    )


def load_data() -> pd.DataFrame:
    if DATA_FILE.exists():
        try:
            df = pd.read_csv(DATA_FILE)
            return df
        except Exception:
            return create_template()
    return create_template()


def save_data(df: pd.DataFrame) -> None:
    df.to_csv(DATA_FILE, index=False, encoding="utf-8-sig")


def prepare_input_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()

    for col in ["วันที่รับยา", "ชื่อยา", "ยอดรับเข้า", "ยอดใช้ไป", "ราคาต่อหน่วย"]:
        if col not in df.columns:
            df[col] = ""

    df["วันที่รับยา"] = df["วันที่รับยา"].astype(str).fillna("").str.strip()
    df["ชื่อยา"] = df["ชื่อยา"].astype(str).fillna("").str.strip()

    df["ยอดรับเข้า"] = pd.to_numeric(df["ยอดรับเข้า"], errors="coerce").fillna(0)
    df["ยอดใช้ไป"] = pd.to_numeric(df["ยอดใช้ไป"], errors="coerce").fillna(0)

    # ราคาต่อหน่วย: แปลงเป็นตัวเลขก่อน แล้วค่อย forward fill ตามชื่อยา
    df["ราคาต่อหน่วย"] = pd.to_numeric(df["ราคาต่อหน่วย"], errors="coerce")

    df = df[df["ชื่อยา"] != ""].copy()

    df["วันที่_dt"] = pd.to_datetime(df["วันที่รับยา"], errors="coerce", dayfirst=True)
    df["row_no"] = range(len(df))

    df = df.sort_values(["ชื่อยา", "วันที่_dt", "row_no"]).reset_index(drop=True)

    # ถ้าแถวไหนไม่ใส่ราคา ให้ใช้ราคาล่าสุดของยาตัวนั้น
    df["ราคาต่อหน่วย"] = (
        df.groupby("ชื่อยา")["ราคาต่อหน่วย"]
        .ffill()
        .fillna(0)
    )

    return df


def calculate_ledger(df: pd.DataFrame) -> pd.DataFrame:
    df = prepare_input_df(df)

    if df.empty:
        return pd.DataFrame(
            columns=[
                "วันที่รับยา",
                "ชื่อยา",
                "ยอดรับเข้า",
                "ยอดใช้ไป",
                "ยอดคงเหลือ",
                "ราคาต่อหน่วย",
                "มูลค่ายาคงเหลือ",
            ]
        )

    df["ยอดรับสะสม"] = df.groupby("ชื่อยา")["ยอดรับเข้า"].cumsum()
    df["ยอดใช้สะสม"] = df.groupby("ชื่อยา")["ยอดใช้ไป"].cumsum()
    df["ยอดคงเหลือ"] = (df["ยอดรับสะสม"] - df["ยอดใช้สะสม"]).clip(lower=0)
    df["มูลค่ายาคงเหลือ"] = df["ยอดคงเหลือ"] * df["ราคาต่อหน่วย"]

    df["วันที่รับยา"] = df["วันที่_dt"].dt.strftime("%d/%m/%y").fillna(df["วันที่รับยา"])

    df = df.drop(columns=["วันที่_dt", "row_no"], errors="ignore")
    return df


def current_summary(ledger: pd.DataFrame) -> pd.DataFrame:
    if ledger.empty:
        return pd.DataFrame(
            columns=[
                "วันที่รับยา",
                "ชื่อยา",
                "ยอดรับสะสม",
                "ยอดใช้สะสม",
                "ยอดคงเหลือ",
                "ราคาต่อหน่วย",
                "มูลค่ายาคงเหลือ",
            ]
        )

    temp = ledger.copy()
    temp["วันที่_dt"] = pd.to_datetime(temp["วันที่รับยา"], errors="coerce", dayfirst=True)
    temp = temp.sort_values(["ชื่อยา", "วันที่_dt"]).reset_index(drop=True)

    summary = temp.groupby("ชื่อยา", as_index=False).tail(1).copy()

    return summary[
        [
            "วันที่รับยา",
            "ชื่อยา",
            "ยอดรับสะสม",
            "ยอดใช้สะสม",
            "ยอดคงเหลือ",
            "ราคาต่อหน่วย",
            "มูลค่ายาคงเหลือ",
        ]
    ].reset_index(drop=True)


def export_excel(summary_df: pd.DataFrame, history_df: pd.DataFrame) -> BytesIO:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        summary_df.to_excel(writer, index=False, sheet_name="Summary")
        history_df.to_excel(writer, index=False, sheet_name="History")
    buffer.seek(0)
    return buffer


def export_word(summary_df: pd.DataFrame, history_df: pd.DataFrame, total_qty: float, total_value: float) -> BytesIO:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("รายงานตัดยอดยาในสถานพยาบาล")
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"ยอดคงเหลือรวม (หน่วย): {total_qty:,.0f}\n")
    p.add_run(f"มูลค่าคงเหลือรวม (บาท): {total_value:,.2f}")

    doc.add_paragraph("")
    doc.add_paragraph("สรุปยอดคงเหลือปัจจุบัน")

    cols = list(summary_df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, col in enumerate(cols):
        table.rows[0].cells[i].text = str(col)

    for _, row in summary_df.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            cells[i].text = str(row[col])

    doc.add_paragraph("")
    doc.add_paragraph("ประวัติการตัดยอดทั้งหมด")

    hcols = list(history_df.columns)
    htable = doc.add_table(rows=1, cols=len(hcols))
    htable.style = "Table Grid"
    for i, col in enumerate(hcols):
        htable.rows[0].cells[i].text = str(col)

    for _, row in history_df.iterrows():
        cells = htable.add_row().cells
        for i, col in enumerate(hcols):
            cells[i].text = str(row[col])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# -------------------------
# โหลดข้อมูลเพื่อแสดงในตาราง
# -------------------------
df = load_data()

display_df = df.copy()
display_df["ราคาต่อหน่วย"] = display_df["ราคาต่อหน่วย"].fillna("").astype(str)
display_df.loc[display_df["ราคาต่อหน่วย"] == "nan", "ราคาต่อหน่วย"] = ""

st.subheader("กรอกข้อมูลยา")

edited_df = st.data_editor(
    display_df,
    num_rows="dynamic",
    use_container_width=True,
    hide_index=True,
    column_config={
        "วันที่รับยา": st.column_config.TextColumn("วันที่รับยา (เช่น 24/5/69)"),
        "ชื่อยา": st.column_config.TextColumn("ชื่อยา"),
        "ยอดรับเข้า": st.column_config.NumberColumn("ยอดรับเข้า", min_value=0, step=1),
        "ยอดใช้ไป": st.column_config.NumberColumn("ยอดใช้ไป", min_value=0, step=1),
        "ราคาต่อหน่วย": st.column_config.TextColumn(
            "ราคาต่อหน่วย",
            help="ใส่ทศนิยมได้ทุกแบบ เช่น 0.05 หรือ 1.237",
        ),
    },
)

# -------------------------
# คำนวณ
# -------------------------
ledger = calculate_ledger(edited_df)
summary = current_summary(ledger)

total_qty = float(summary["ยอดคงเหลือ"].sum()) if not summary.empty else 0.0
total_value = float(summary["มูลค่ายาคงเหลือ"].sum()) if not summary.empty else 0.0

st.divider()

st.subheader("สรุปยอดคงเหลือปัจจุบัน")
c1, c2 = st.columns(2)
c1.metric("ยอดคงเหลือรวม (หน่วย)", f"{total_qty:,.0f}")
c2.metric("มูลค่าคงเหลือรวม (บาท)", f"{total_value:,.2f}")

st.subheader("ตารางสรุปผลปัจจุบัน")
st.dataframe(summary, use_container_width=True, hide_index=True)

st.subheader("ประวัติการตัดยอดทั้งหมด")
st.dataframe(
    ledger[
        [
            "วันที่รับยา",
            "ชื่อยา",
            "ยอดรับเข้า",
            "ยอดใช้ไป",
            "ยอดคงเหลือ",
            "ราคาต่อหน่วย",
            "มูลค่ายาคงเหลือ",
        ]
    ],
    use_container_width=True,
    hide_index=True,
)

# -------------------------
# ปุ่มบันทึก / ดาวน์โหลด
# -------------------------
col1, col2, col3, col4 = st.columns(4)

with col1:
    if st.button("บันทึกข้อมูลลง CSV"):
        save_data(edited_df)
        st.success("บันทึกข้อมูลเรียบร้อยแล้ว")

with col2:
    csv_data = ledger.to_csv(index=False, encoding="utf-8-sig")
    st.download_button(
        label="ดาวน์โหลด CSV",
        data=csv_data,
        file_name="medicine_stock_result.csv",
        mime="text/csv",
    )

with col3:
    excel_buffer = export_excel(summary, ledger)
    st.download_button(
        label="ดาวน์โหลด Excel",
        data=excel_buffer,
        file_name="medicine_stock_report.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

with col4:
    word_buffer = export_word(summary, ledger, total_qty, total_value)
    st.download_button(
        label="ดาวน์โหลด Word",
        data=word_buffer,
        file_name="medicine_stock_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
